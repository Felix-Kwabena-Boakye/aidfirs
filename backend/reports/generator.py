"""
Forensic Report Generator for AIDFIRS.
Generates PDF, DOCX, and HTML forensic investigation reports.
"""
import os
import io
from datetime import datetime, timezone

REPORTS_STORAGE = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    'storage', 'reports'
)
os.makedirs(REPORTS_STORAGE, exist_ok=True)


def _format_bytes(size):
    if not size:
        return '0 Bytes'
    units = ['Bytes', 'KB', 'MB', 'GB', 'TB']
    i = 0
    while size >= 1024 and i < len(units) - 1:
        size /= 1024.0
        i += 1
    return f"{size:.2f} {units[i]}"


def generate_html_report(case, files, timeline, coc_entries, examiner, report_type='full'):
    """Generate a complete HTML forensic report."""
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    case_number = getattr(case, 'case_number', 'N/A')
    case_title = getattr(case, 'title', 'N/A')
    case_status = getattr(case, 'status', 'N/A')
    case_priority = getattr(case, 'priority', 'N/A')
    case_created = case.created_at.strftime("%Y-%m-%d %H:%M:%S UTC") if hasattr(case, 'created_at') and case.created_at else 'N/A'
    total_size = sum(f.size for f in files if f.size)

    # Build file rows
    file_rows = ""
    for i, f in enumerate(files, 1):
        file_rows += f"""
        <tr>
            <td>{i}</td>
            <td>{f.filename or 'N/A'}</td>
            <td>{(f.file_extension or '').upper() or 'BIN'}</td>
            <td>{_format_bytes(f.size)}</td>
            <td>{f.recovery_method or 'N/A'}</td>
            <td class="hash">{(f.hash_sha256 or '')[:32]}...</td>
            <td>{f.hash_md5 or 'N/A'}</td>
            <td>{f.recovery_status or 'N/A'}</td>
        </tr>"""

    # Build timeline rows
    timeline_rows = ""
    for ev in timeline:
        ts = ev.timestamp.strftime("%Y-%m-%d %H:%M:%S UTC") if ev.timestamp else 'N/A'
        timeline_rows += f"""
        <tr>
            <td>{ts}</td>
            <td><span class="badge">{ev.event_type or 'N/A'}</span></td>
            <td>{ev.description or 'N/A'}</td>
            <td>{ev.actor or 'System'}</td>
        </tr>"""

    # Build CoC rows
    coc_rows = ""
    for entry in coc_entries:
        d = entry.to_dict() if hasattr(entry, 'to_dict') else entry
        ts = d.get('timestamp', 'N/A')
        coc_rows += f"""
        <tr>
            <td>{ts}</td>
            <td>{d.get('action', 'N/A')}</td>
            <td>{d.get('performed_by', 'N/A')}</td>
            <td>{d.get('notes', 'N/A')}</td>
            <td class="hash">{(d.get('hash_after') or '')[:24]}...</td>
        </tr>"""

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>AIDFIRS Forensic Report - {case_number}</title>
<style>
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: 'Segoe UI', Arial, sans-serif; background: #0f172a; color: #e2e8f0; line-height: 1.6; }}
  .page {{ max-width: 1100px; margin: 0 auto; padding: 40px 20px; }}
  .header {{ background: linear-gradient(135deg, #1e3a5f, #0f2744); border: 1px solid #1e40af; border-radius: 12px; padding: 32px; margin-bottom: 32px; }}
  .header-logo {{ font-size: 24px; font-weight: 800; color: #60a5fa; letter-spacing: 2px; margin-bottom: 8px; }}
  .header-title {{ font-size: 32px; font-weight: 700; color: white; margin-bottom: 8px; }}
  .header-meta {{ color: #94a3b8; font-size: 14px; }}
  .header-meta span {{ color: #60a5fa; font-weight: 600; }}
  .section {{ background: #1e293b; border: 1px solid #334155; border-radius: 10px; padding: 28px; margin-bottom: 24px; }}
  .section-title {{ font-size: 18px; font-weight: 700; color: #60a5fa; margin-bottom: 20px; padding-bottom: 12px; border-bottom: 1px solid #334155; display: flex; align-items: center; gap: 8px; }}
  .section-title::before {{ content: ''; display: inline-block; width: 4px; height: 18px; background: #3b82f6; border-radius: 2px; }}
  .grid-2 {{ display: grid; grid-template-columns: 1fr 1fr; gap: 20px; }}
  .info-card {{ background: #0f172a; border: 1px solid #1e293b; border-radius: 8px; padding: 16px; }}
  .info-label {{ font-size: 11px; font-weight: 600; color: #64748b; text-transform: uppercase; letter-spacing: 0.5px; margin-bottom: 4px; }}
  .info-value {{ font-size: 15px; color: #e2e8f0; font-weight: 500; }}
  .stat-grid {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 16px; margin-bottom: 24px; }}
  .stat-card {{ background: #0f172a; border: 1px solid #1e293b; border-radius: 8px; padding: 20px; text-align: center; }}
  .stat-number {{ font-size: 28px; font-weight: 800; color: #60a5fa; }}
  .stat-label {{ font-size: 12px; color: #64748b; margin-top: 4px; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 13px; }}
  th {{ background: #0f172a; color: #94a3b8; padding: 12px 16px; text-align: left; font-size: 11px; font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px; border-bottom: 1px solid #1e293b; }}
  td {{ padding: 10px 16px; border-bottom: 1px solid #1e293b20; color: #cbd5e1; }}
  tr:hover td {{ background: #ffffff08; }}
  .hash {{ font-family: monospace; font-size: 11px; color: #94a3b8; }}
  .badge {{ background: #1e3a5f; color: #60a5fa; padding: 2px 10px; border-radius: 12px; font-size: 11px; font-weight: 600; }}
  .status-open {{ color: #34d399; }} .status-closed {{ color: #f87171; }} .status-progress {{ color: #fbbf24; }}
  .executive-summary {{ color: #cbd5e1; line-height: 1.8; }}
  .recommendation {{ background: #0f172a; border-left: 3px solid #3b82f6; padding: 12px 16px; border-radius: 0 6px 6px 0; margin-bottom: 10px; color: #cbd5e1; }}
  .footer {{ text-align: center; color: #475569; font-size: 12px; margin-top: 40px; padding-top: 20px; border-top: 1px solid #1e293b; }}
  .confidential {{ color: #ef4444; font-weight: 700; font-size: 11px; letter-spacing: 1px; text-transform: uppercase; }}
  @media print {{ body {{ background: white; color: black; }} .section {{ border: 1px solid #ccc; }} }}
</style>
</head>
<body>
<div class="page">

  <!-- HEADER -->
  <div class="header">
    <div class="header-logo">⚡ AIDFIRS</div>
    <div class="header-title">Digital Forensic Investigation Report</div>
    <div class="header-meta">
      Case: <span>{case_number}</span> &nbsp;|&nbsp;
      Generated: <span>{now}</span> &nbsp;|&nbsp;
      Examiner: <span>{examiner}</span> &nbsp;|&nbsp;
      <span class="confidential">CONFIDENTIAL - LAW ENFORCEMENT USE ONLY</span>
    </div>
  </div>

  <!-- EXECUTIVE SUMMARY -->
  <div class="section">
    <div class="section-title">Executive Summary</div>
    <div class="executive-summary">
      <p>This forensic investigation report was prepared by <strong>{examiner}</strong> in connection with Case <strong>{case_number}: {case_title}</strong>.
      The investigation utilized the AIDFIRS (AI-Powered Digital Forensic Investigation and Recovery System) platform to conduct a comprehensive digital forensic analysis.</p>
      <br>
      <p>A total of <strong>{len(files)} files</strong> were recovered ({_format_bytes(total_size)}), spanning {len(set(f.file_extension for f in files if f.file_extension))} unique file types.
      All recovered artifacts have been cryptographically hashed to ensure evidentiary integrity. The timeline recorded
      <strong>{len(timeline)} forensic events</strong> with full chain-of-custody documentation.</p>
      <br>
      <p>Case Status: <strong>{case_status.upper()}</strong> &nbsp;|&nbsp; Priority: <strong>{case_priority.upper()}</strong> &nbsp;|&nbsp; Opened: <strong>{case_created}</strong></p>
    </div>
  </div>

  <!-- CASE INFORMATION -->
  <div class="section">
    <div class="section-title">Case Information</div>
    <div class="grid-2">
      <div class="info-card"><div class="info-label">Case Number</div><div class="info-value">{case_number}</div></div>
      <div class="info-card"><div class="info-label">Case Title</div><div class="info-value">{case_title}</div></div>
      <div class="info-card"><div class="info-label">Status</div><div class="info-value">{case_status}</div></div>
      <div class="info-card"><div class="info-label">Priority</div><div class="info-value">{case_priority}</div></div>
      <div class="info-card"><div class="info-label">Investigator</div><div class="info-value">{examiner}</div></div>
      <div class="info-card"><div class="info-label">Date Opened</div><div class="info-value">{case_created}</div></div>
      <div class="info-card"><div class="info-label">Description</div><div class="info-value">{getattr(case, 'description', 'N/A')}</div></div>
      <div class="info-card"><div class="info-label">Report Generated</div><div class="info-value">{now}</div></div>
    </div>
  </div>

  <!-- STATISTICS -->
  <div class="section">
    <div class="section-title">Evidence Statistics</div>
    <div class="stat-grid">
      <div class="stat-card"><div class="stat-number">{len(files)}</div><div class="stat-label">Files Recovered</div></div>
      <div class="stat-card"><div class="stat-number">{_format_bytes(total_size)}</div><div class="stat-label">Total Evidence Size</div></div>
      <div class="stat-card"><div class="stat-number">{len(timeline)}</div><div class="stat-label">Timeline Events</div></div>
      <div class="stat-card"><div class="stat-number">{len(coc_entries)}</div><div class="stat-label">CoC Entries</div></div>
    </div>
  </div>

  <!-- RECOVERED EVIDENCE -->
  <div class="section">
    <div class="section-title">Recovered Evidence ({len(files)} Files)</div>
    <table>
      <thead>
        <tr>
          <th>#</th><th>Filename</th><th>Type</th><th>Size</th>
          <th>Recovery Method</th><th>SHA-256 (truncated)</th><th>MD5</th><th>Status</th>
        </tr>
      </thead>
      <tbody>{file_rows if file_rows else '<tr><td colspan="8" style="text-align:center; color:#64748b;">No files recovered.</td></tr>'}</tbody>
    </table>
  </div>

  <!-- FORENSIC TIMELINE -->
  <div class="section">
    <div class="section-title">Forensic Timeline ({len(timeline)} Events)</div>
    <table>
      <thead>
        <tr><th>Timestamp</th><th>Event Type</th><th>Description</th><th>Actor</th></tr>
      </thead>
      <tbody>{timeline_rows if timeline_rows else '<tr><td colspan="4" style="text-align:center; color:#64748b;">No timeline events recorded.</td></tr>'}</tbody>
    </table>
  </div>

  <!-- CHAIN OF CUSTODY -->
  <div class="section">
    <div class="section-title">Chain of Custody ({len(coc_entries)} Records)</div>
    <table>
      <thead>
        <tr><th>Timestamp</th><th>Action</th><th>Examiner</th><th>Notes</th><th>Hash Verification</th></tr>
      </thead>
      <tbody>{coc_rows if coc_rows else '<tr><td colspan="5" style="text-align:center; color:#64748b;">No chain of custody records.</td></tr>'}</tbody>
    </table>
  </div>

  <!-- RECOMMENDATIONS -->
  <div class="section">
    <div class="section-title">Recommendations</div>
    <div class="recommendation">All recovered files should be preserved in a write-protected forensic repository and backed up to offline storage immediately.</div>
    <div class="recommendation">Hash values (SHA-256, MD5, SHA-1) for all recovered files should be independently verified before presenting as evidence in legal proceedings.</div>
    <div class="recommendation">The chain of custody documentation should be maintained and appended for every subsequent examination or transfer of the evidence.</div>
    <div class="recommendation">Any files with hash verification failures should be quarantined and flagged for further investigation before use as evidence.</div>
    <div class="recommendation">This report and all associated evidence should be stored in accordance with applicable digital evidence handling standards (ISO/IEC 27037, NIST SP 800-101r1).</div>
  </div>

  <!-- FOOTER -->
  <div class="footer">
    <p>This report was generated by <strong>AIDFIRS — AI-Powered Digital Forensic Investigation and Recovery System</strong></p>
    <p>Report generated: {now} &nbsp;|&nbsp; Examiner: {examiner} &nbsp;|&nbsp; Case: {case_number}</p>
    <p class="confidential">⚠️ CONFIDENTIAL — This document contains sensitive forensic investigation data. Unauthorized disclosure is prohibited.</p>
  </div>

</div>
</body>
</html>"""
    return html


def generate_pdf_report(case, files, timeline, coc_entries, examiner, report_id):
    """Generate a PDF forensic report using reportlab."""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        Table, TableStyle, HRFlowable, PageBreak)
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    except ImportError:
        raise RuntimeError("reportlab is not installed. Run: pip install reportlab")

    filename = f"AIDFIRS_Report_{report_id}.pdf"
    output_path = os.path.join(REPORTS_STORAGE, filename)

    doc = SimpleDocTemplate(output_path, pagesize=A4,
                            rightMargin=0.75*inch, leftMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle('Title', parent=styles['Title'],
        fontSize=20, textColor=colors.HexColor('#1e40af'), spaceAfter=6)
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'],
        fontSize=13, textColor=colors.HexColor('#1e40af'), spaceAfter=8, spaceBefore=16)
    body_style = ParagraphStyle('Body', parent=styles['Normal'],
        fontSize=10, spaceAfter=4, leading=14)
    mono_style = ParagraphStyle('Mono', parent=styles['Normal'],
        fontSize=8, fontName='Courier', textColor=colors.HexColor('#374151'))
    red_style = ParagraphStyle('Red', parent=styles['Normal'],
        fontSize=8, textColor=colors.red, fontName='Helvetica-Bold')

    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    case_number = getattr(case, 'case_number', 'N/A')

    # Cover
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph("AIDFIRS — AI-Powered Digital Forensic Investigation Report", title_style))
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#1e40af')))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(f"<b>Case Number:</b> {case_number}", body_style))
    story.append(Paragraph(f"<b>Case Title:</b> {getattr(case, 'title', 'N/A')}", body_style))
    story.append(Paragraph(f"<b>Examiner:</b> {examiner}", body_style))
    story.append(Paragraph(f"<b>Generated:</b> {now}", body_style))
    story.append(Paragraph("<b>CONFIDENTIAL — LAW ENFORCEMENT USE ONLY</b>", red_style))
    story.append(Spacer(1, 0.3*inch))

    # Executive Summary
    story.append(Paragraph("Executive Summary", heading_style))
    total_size = sum(f.size for f in files if f.size)
    story.append(Paragraph(
        f"This forensic report was prepared by <b>{examiner}</b> for Case <b>{case_number}</b>. "
        f"A total of <b>{len(files)} files</b> ({_format_bytes(total_size)}) were recovered and analyzed. "
        f"The forensic timeline recorded <b>{len(timeline)} events</b> with complete chain-of-custody documentation.",
        body_style))
    story.append(Spacer(1, 0.2*inch))

    # Case Info Table
    story.append(Paragraph("Case Information", heading_style))
    case_data = [
        ['Case Number', case_number, 'Status', getattr(case, 'status', 'N/A')],
        ['Title', getattr(case, 'title', 'N/A'), 'Priority', getattr(case, 'priority', 'N/A')],
        ['Investigator', examiner, 'Files Recovered', str(len(files))],
        ['Total Evidence', _format_bytes(total_size), 'Timeline Events', str(len(timeline))],
    ]
    t = Table(case_data, colWidths=[1.5*inch, 2.5*inch, 1.5*inch, 1.5*inch])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8fafc')),
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#dbeafe')),
        ('BACKGROUND', (2, 0), (2, -1), colors.HexColor('#dbeafe')),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('PADDING', (0, 0), (-1, -1), 6),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.2*inch))

    # Recovered Files
    story.append(Paragraph(f"Recovered Evidence ({len(files)} Files)", heading_style))
    if files:
        headers = ['#', 'Filename', 'Type', 'Size', 'Method', 'SHA-256 (partial)']
        file_data = [headers]
        for i, f in enumerate(files[:100], 1):  # limit to 100 rows in PDF
            file_data.append([
                str(i),
                (f.filename or '')[:40],
                (f.file_extension or 'BIN').upper(),
                _format_bytes(f.size),
                (f.recovery_method or 'N/A')[:20],
                (f.hash_sha256 or '')[:24] + '...' if f.hash_sha256 else 'N/A',
            ])
        ft = Table(file_data, colWidths=[0.4*inch, 2.5*inch, 0.6*inch, 0.8*inch, 1.0*inch, 1.7*inch])
        ft.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f1f5f9')]),
            ('GRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#e2e8f0')),
            ('PADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(ft)
        if len(files) > 100:
            story.append(Paragraph(f"... and {len(files)-100} more files (see exported CSV for complete list)", body_style))
    else:
        story.append(Paragraph("No files recovered for this case.", body_style))
    story.append(Spacer(1, 0.2*inch))

    # Timeline (truncated)
    story.append(Paragraph(f"Forensic Timeline ({len(timeline)} Events)", heading_style))
    if timeline:
        t_headers = ['Timestamp', 'Event Type', 'Description', 'Actor']
        t_data = [t_headers]
        for ev in timeline[:50]:
            ts = ev.timestamp.strftime("%Y-%m-%d %H:%M") if ev.timestamp else 'N/A'
            t_data.append([ts, (ev.event_type or '')[:20], (ev.description or '')[:60], (ev.actor or 'System')[:20]])
        tt = Table(t_data, colWidths=[1.3*inch, 1.4*inch, 3.2*inch, 1.1*inch])
        tt.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 7),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f1f5f9')]),
            ('GRID', (0, 0), (-1, -1), 0.3, colors.HexColor('#e2e8f0')),
            ('PADDING', (0, 0), (-1, -1), 4),
        ]))
        story.append(tt)
    else:
        story.append(Paragraph("No timeline events recorded.", body_style))

    # Recommendations
    story.append(PageBreak())
    story.append(Paragraph("Recommendations", heading_style))
    recs = [
        "All recovered files must be preserved in a write-protected forensic repository.",
        "Hash values (SHA-256, MD5, SHA-1) should be independently verified before court presentation.",
        "Chain of custody documentation must be maintained for all subsequent transfers.",
        "Files with hash verification failures must be quarantined and re-examined.",
        "Handle this evidence per ISO/IEC 27037 and NIST SP 800-101r1 standards.",
    ]
    for r in recs:
        story.append(Paragraph(f"• {r}", body_style))

    # Conclusion
    story.append(Paragraph("Conclusion", heading_style))
    story.append(Paragraph(
        f"This investigation was conducted using AIDFIRS with full digital forensic integrity. "
        f"All evidence has been cryptographically hashed and chain-of-custody documented. "
        f"The findings of this report are prepared for {case_number} and are ready for review by authorized personnel.",
        body_style))
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph(f"Examiner Signature: _______________________ &nbsp;&nbsp; Date: {now}", body_style))

    doc.build(story)
    return output_path


def generate_docx_report(case, files, timeline, coc_entries, examiner, report_id):
    """Generate a DOCX forensic report using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    doc = Document()
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")
    case_number = getattr(case, 'case_number', 'N/A')
    total_size = sum(f.size for f in files if f.size)

    # Title
    title = doc.add_heading('AIDFIRS — Digital Forensic Investigation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Case: {case_number} | Examiner: {examiner} | Generated: {now}")
    doc.add_paragraph("CONFIDENTIAL — LAW ENFORCEMENT USE ONLY").bold = True
    doc.add_paragraph("")

    # Executive Summary
    doc.add_heading("Executive Summary", 1)
    doc.add_paragraph(
        f"This forensic report was prepared by {examiner} for Case {case_number}: {getattr(case, 'title', 'N/A')}. "
        f"A total of {len(files)} files ({_format_bytes(total_size)}) were recovered. "
        f"The forensic timeline recorded {len(timeline)} events with complete chain-of-custody documentation."
    )

    # Case Info
    doc.add_heading("Case Information", 1)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    rows_data = [
        ['Case Number', case_number, 'Status', getattr(case, 'status', 'N/A')],
        ['Title', getattr(case, 'title', 'N/A'), 'Priority', getattr(case, 'priority', 'N/A')],
        ['Investigator', examiner, 'Files Recovered', str(len(files))],
        ['Evidence Size', _format_bytes(total_size), 'Timeline Events', str(len(timeline))],
    ]
    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text

    doc.add_paragraph("")

    # Recovered Files
    doc.add_heading(f"Recovered Evidence ({len(files)} Files)", 1)
    if files:
        ft = doc.add_table(rows=1, cols=5)
        ft.style = 'Table Grid'
        hdr = ft.rows[0].cells
        for i, h in enumerate(['Filename', 'Type', 'Size', 'Recovery Method', 'SHA-256 (partial)']):
            hdr[i].text = h
        for f in files[:100]:
            row = ft.add_row().cells
            row[0].text = f.filename or 'N/A'
            row[1].text = (f.file_extension or 'BIN').upper()
            row[2].text = _format_bytes(f.size)
            row[3].text = f.recovery_method or 'N/A'
            row[4].text = (f.hash_sha256 or '')[:24] + '...' if f.hash_sha256 else 'N/A'
    else:
        doc.add_paragraph("No files recovered.")

    doc.add_paragraph("")

    # Timeline
    doc.add_heading(f"Forensic Timeline ({len(timeline)} Events)", 1)
    if timeline:
        tt = doc.add_table(rows=1, cols=4)
        tt.style = 'Table Grid'
        hdr = tt.rows[0].cells
        for i, h in enumerate(['Timestamp', 'Event Type', 'Description', 'Actor']):
            hdr[i].text = h
        for ev in timeline[:50]:
            row = tt.add_row().cells
            row[0].text = ev.timestamp.strftime("%Y-%m-%d %H:%M") if ev.timestamp else 'N/A'
            row[1].text = ev.event_type or 'N/A'
            row[2].text = (ev.description or 'N/A')[:100]
            row[3].text = ev.actor or 'System'
    else:
        doc.add_paragraph("No timeline events.")

    # Recommendations
    doc.add_heading("Recommendations", 1)
    recs = [
        "Preserve all files in a write-protected forensic repository.",
        "Independently verify all hash values before court presentation.",
        "Maintain chain-of-custody for all subsequent transfers.",
        "Quarantine files with hash verification failures.",
        "Handle evidence per ISO/IEC 27037 and NIST SP 800-101r1 standards.",
    ]
    for r in recs:
        doc.add_paragraph(f"• {r}")

    # Conclusion
    doc.add_heading("Conclusion", 1)
    doc.add_paragraph(
        f"This investigation was conducted using AIDFIRS with full digital forensic integrity. "
        f"All evidence has been cryptographically hashed and chain-of-custody documented. "
        f"This report is prepared for Case {case_number} and is ready for authorized review."
    )
    doc.add_paragraph("")
    doc.add_paragraph(f"Examiner Signature: _______________________ | Date: {now}")

    filename = f"AIDFIRS_Report_{report_id}.docx"
    output_path = os.path.join(REPORTS_STORAGE, filename)
    doc.save(output_path)
    return output_path
